"""Generate Word report and script - two separate files"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

GREEN = RGBColor(0x2E, 0x7D, 0x32)
GREEN_L = RGBColor(0x4C, 0xAF, 0x50)
GREEN_D = RGBColor(0x1B, 0x5E, 0x20)
GREEN_P = RGBColor(0xE8, 0xF5, 0xE9)
GREEN_MUTED = RGBColor(0x66, 0xBB, 0x6A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
DARK = RGBColor(0x1A, 0x1A, 0x1A)
GRAY = RGBColor(0x88, 0x88, 0x88)
LGRAY = RGBColor(0xF5, 0xF5, 0xF5)

# ═══════════════════════════════════════
# COMMON SETUP
# ═══════════════════════════════════════
def setup_doc():
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin = Cm(2.8)
        sec.right_margin = Cm(2.8)
    # Normal style
    normal = doc.styles['Normal']
    normal.font.name = '微软雅黑'
    normal.font.size = Pt(11)
    normal.font.color.rgb = DARK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.5
    rPr = normal.element.get_or_add_rPr()
    rFonts = rPr.makeelement(qn('w:rFonts'), {})
    rFonts.set(qn('w:eastAsia'), '微软雅黑')
    rFonts.set(qn('w:ascii'), '微软雅黑')
    rFonts.set(qn('w:hAnsi'), '微软雅黑')
    rPr.insert(0, rFonts)
    # Heading 1
    h1s = doc.styles['Heading 1']
    h1s.font.name = '微软雅黑'; h1s.font.size = Pt(18); h1s.font.bold = True
    h1s.font.color.rgb = GREEN
    h1s.paragraph_format.space_before = Pt(24)
    h1s.paragraph_format.space_after = Pt(12)
    h1s.paragraph_format.first_line_indent = Cm(0)
    hPr = h1s.element.get_or_add_rPr()
    hRf = hPr.makeelement(qn('w:rFonts'), {})
    hRf.set(qn('w:eastAsia'), '微软雅黑'); hPr.insert(0, hRf)
    # Heading 2
    h2s = doc.styles['Heading 2']
    h2s.font.name = '微软雅黑'; h2s.font.size = Pt(13); h2s.font.bold = True
    h2s.font.color.rgb = GREEN
    h2s.paragraph_format.space_before = Pt(16)
    h2s.paragraph_format.space_after = Pt(8)
    h2s.paragraph_format.first_line_indent = Cm(0)
    hPr2 = h2s.element.get_or_add_rPr()
    hRf2 = hPr2.makeelement(qn('w:rFonts'), {})
    hRf2.set(qn('w:eastAsia'), '微软雅黑'); hPr2.insert(0, hRf2)
    return doc

def p(doc, text, bold=False, size=11, center=False, indent=False, color=DARK):
    para = doc.add_paragraph()
    if center: para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if indent: para.paragraph_format.first_line_indent = Cm(0.7)
    r = para.add_run(text)
    if size: r.font.size = Pt(size)
    r.bold = bold
    r.font.color.rgb = color
    return para

def h1(doc, text):
    return doc.add_heading(text, level=1)

def h2(doc, text):
    return doc.add_heading(text, level=2)

def dot(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(1.2)
    para.paragraph_format.first_line_indent = Cm(-0.5)
    r = para.add_run('· ' + text)
    return para

def sep(doc):
    pr = doc.add_paragraph()
    pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = pr.add_run('━' * 40)
    rr.font.size = Pt(8)
    rr.font.color.rgb = GREEN_L

# ═══════════════════════════════════════
# FILE 1: REPORT
# ═══════════════════════════════════════
doc = setup_doc()

# COVER
for _ in range(7): doc.add_paragraph()
p(doc, '拾  句', bold=True, size=34, center=True, color=GREEN)
p(doc, '一词入 · 千年应', size=15, center=True, color=RGBColor(0x88,0x77,0x66))
for _ in range(4): doc.add_paragraph()
p(doc, 'AI 智能体开发 · 期末作业报告', size=13, center=True, color=GRAY)
p(doc, '2025-2026学年第二学期', size=11, center=True, color=GRAY)
for _ in range(2): doc.add_paragraph()
p(doc, '王  宸', size=18, bold=True, center=True)
p(doc, '2025212714', size=12, center=True, color=GREEN)
doc.add_page_break()

# ── 一 ──
h1(doc, '一、制作动机')
p(doc, '每个人都有过这样的时刻——心里堵着一团情绪，却找不到合适的语言来表达。想发条朋友圈，打了又删，总觉得怎么说都差一点。', indent=True)
p(doc, '但翻开唐诗宋词，你会发现——"孤独"早就被陈子昂写透了："前不见古人，后不见来者，念天地之悠悠，独怆然而涕下。"思念被李清照凝成了"此情无计可消除，才下眉头，却上心头。"后悔被归有光埋在了一棵枇杷树里。', indent=True)
p(doc, '明明古人早就替我们写尽了人间悲欢，但我们和那些句子之间隔着一千年。不是那些句子不好，是我们根本不知道它们在哪。', indent=True)
p(doc, '市面上有无数的AI工具，但它们做的事情都一样——你问它答。你让它"写首诗"，它给你几行押韵的字，读起来像流水线产品。真正动人的诗句不需要AI来写——李白杜甫早就写好了，你只是需要一个懂你的人帮你找到它。', indent=True)
p(doc, '这就是「拾句」的出发点：不是让AI写诗，而是让AI替你翻遍唐诗宋词、古文散文，找到那句刚好砸中你心口的古典名句。输入一个词，遇见一千年前同样有过这种心情的人。', indent=True, bold=True)

# ── 二 ──
h1(doc, '二、制作中遇到的问题和困难')
h2(doc, '问题1：如何让AI精准匹配古典文学，而不是胡编乱造？')
p(doc, '让AI写诗很容易——但写出来的是三流作品，毫无文学价值。真正有价值的是引用和化用古典名句。但大语言模型存在"幻觉"问题：你问它"和月亮有关的唐诗"，它可能编一句听起来像唐诗但根本不存在的句子。', indent=True)
p(doc, '解决方案：', indent=True, bold=True)
dot(doc, '将古典文学分为四个明确的赛道——唐诗（李白、杜甫、王维等）、宋词（苏轼、李清照、辛弃疾等）、古文（归有光、张岱、沈复等）、散文（朱自清、汪曾祺、余光中等）。每个赛道在Prompt中明确指定作者范围和文风特征。')
dot(doc, '要求AI输出结构化JSON格式 {quote, source, paraphrase, emotion}，其中source字段强制要求标注"作者·篇名"，倒逼AI确认来源，减少幻觉。')
dot(doc, '设置temperature=0.9保持输出多样性，同时通过详细的System Prompt约束质量底线。')

h2(doc, '问题2：如何让古典名句和现代人的情绪产生共鸣？')
p(doc, '如果只是关键词机械匹配（输入"月亮"就输出"床前明月光"），太浅了，没有任何惊喜感。真正动人的匹配是意境层面的——比如输入"失眠"，匹配到的应该是"辗转反侧""夜不能寐"这种情感层面的连接，而非含"睡"字的句子。', indent=True)
p(doc, '解决方案：', indent=True, bold=True)
dot(doc, 'Prompt中强调"化用"而非"搜索"。AI先理解关键词的情感内核（孤独→苍茫/辽阔/一人），再从情感角度匹配古典名句。')
dot(doc, '要求输出一句白话"意境解读"（paraphrase），在古典原文和现代感受之间架一座桥。')

h2(doc, '问题3：部署后海外IP被DeepSeek拒绝')
p(doc, '项目部署在Streamlit Cloud（服务器位于美国），DeepSeek官方API会拒绝来自海外IP的请求。首次部署后所有API调用均返回403错误，应用完全无法使用。', indent=True)
p(doc, '解决方案：', indent=True, bold=True)
dot(doc, '切换到国内API代理平台「硅基流动（SiliconFlow）」，底层模型仍为DeepSeek-V3，但API服务器在国内，不受IP地域限制。')
dot(doc, 'API Key通过Streamlit Cloud的Secrets机制管理，不在代码中硬编码，保障安全性。')
dot(doc, '增加60秒超时和try/catch错误处理，API调用失败时显示具体错误信息而非白屏。')

h2(doc, '问题4：HTML卡片渲染出现乱码')
p(doc, '卡片使用CSS+HTML在Streamlit中渲染。初版使用Python多行字符串拼接构建HTML，CSS渐变中的特殊字符导致HTML结构断裂，卡片显示为源代码文本。', indent=True)
p(doc, '解决方案：', indent=True, bold=True)
dot(doc, '改用Python f-string构建HTML字符串，单行输出避免换行符干扰。')
dot(doc, '对所有AI生成的文本内容使用html.escape()进行转义处理，防止特殊字符破坏HTML结构。')
dot(doc, '装饰元素使用独立CSS类而非内联复杂样式，提高可维护性。')

# ── 三 ──
h1(doc, '三、技术架构')
p(doc, '技术栈总览：', indent=True, bold=True)
dot(doc, '开发语言：Python 3.x')
dot(doc, '前端框架：Streamlit（纯Python Web UI）')
dot(doc, '大语言模型：DeepSeek-V3（通过硅基流动API调用，兼容OpenAI SDK格式）')
dot(doc, '部署平台：Streamlit Cloud（免费层，与GitHub仓库自动同步部署）')
dot(doc, '代码托管：GitHub Public仓库（ZhiRam/AI-agent-work）')
p(doc, '')
p(doc, '核心调用链：', bold=True)
p(doc, '用户输入关键字 → Streamlit UI → agent.generate_card() → llm_client.chat() → OpenAI SDK → SiliconFlow API（DeepSeek-V3）→ 返回JSON → render_card() → HTML/CSS卡片', indent=True)
p(doc, '')
p(doc, '项目文件结构与职责：', bold=True)
dot(doc, 'app.py —— Streamlit主界面、CSS样式系统、卡片渲染函数render_card()')
dot(doc, 'agent.py —— LLM调用逻辑、JSON响应解析、卡片内容生成')
dot(doc, 'prompts.py —— 四种风格的System Prompt定义、CSS配色方案')
dot(doc, 'llm_client.py —— OpenAI SDK封装层，含超时与异常处理')
dot(doc, 'config.py —— API地址、模型名称、温度参数等全局配置')
p(doc, '')
p(doc, '四种风格的技术实现：', bold=True)
p(doc, '每种风格在STYLES字典中以键值对完整定义，包含CSS变量（背景渐变、文字颜色、强调色、字体栈）和文学Prompt。render_card()运行时根据风格动态生成HTML，纯CSS实现所有装饰效果，无需外部图片资源。', indent=True)

# ── 四 ──
h1(doc, '四、创新点')
h2(doc, '创新点1：AI不写诗，AI帮你找诗')
p(doc, '不同于市面上所有"AI作诗"类产品，本项目拒绝让AI创作内容，而是让AI做内容检索和意境匹配。利用大语言模型训练数据中内化的古典文学知识，在语义层面连接"现代关键词"和"古典名句"。本质是有温度的搜索，而非无灵魂的创作。大多数AI产品追求"生成更多"，本项目追求"筛选最好的"。', indent=True)
h2(doc, '创新点2：一词四境，横跨千年文学史')
p(doc, '同一个关键词在唐诗、宋词、古文、散文四种体裁中找到完全不同的美学表达。例如"孤独"——唐诗有"前不见古人，后不见来者"的苍茫辽阔；宋词有"无可奈何花落去"的婉约低徊；古文有"庭有枇杷树，今已亭亭如盖矣"的深情克制；散文有"热闹是他们的，我什么也没有"的现代孤独。用户切换风格即可在四种文学传统之间穿行，同一情绪获得四种不同的精神回应。', indent=True)
h2(doc, '创新点3：极简交互，零学习成本')
p(doc, '整个界面只有三个可操作元素：输入框、风格按钮、生成按钮。没有任何设置项，不需要写Prompt，不需要调参数。3秒上手，1秒出结果。卡片设计精美可直接截图分享——产品即内容。', indent=True)
h2(doc, '创新点4：视觉设计服务于文学内容')
p(doc, '不同文学风格对应完全不同的视觉语言——唐诗采用宣纸底色、楷体字、红色印章，营造盛唐气象；宋词采用淡紫渐变、柔光圆环，呼应婉约意境；古文采用古纸色调、邮戳印记，呈现文人雅致；散文采用暖黄底色、衬线字体，保留人间烟火温度。视觉不是装饰层，而是文学体验的组成部分。', indent=True)

# ── 五 ──
h1(doc, '五、缺点与反思')
h2(doc, '当前不足')
p(doc, '1. 匹配精度受限于模型训练数据的古典文学覆盖度。DeepSeek-V3在部分偏门文献上存在不足，偶尔出现匹配不贴切的情况。此外模型偶尔会将作者搞错，目前缺乏校验机制。', indent=True)
p(doc, '2. 缺乏用户反馈闭环。系统无法获知用户对匹配结果的满意度，缺少点赞机制来积累数据以持续优化匹配策略。', indent=True)
p(doc, '3. 移动端适配不足。Streamlit默认布局在手机上体验不佳，卡片固定宽度导致小屏幕需要横向滚动。', indent=True)
p(doc, '4. 风格数量有限。目前仅四种风格，但扩展空间很大——诗经的古朴、元曲的俚俗、现代诗的先锋，每种只需在STYLES字典中添加一个条目。', indent=True)
p(doc, '5. 缺少个性化记忆。每次使用是独立会话，系统不记得用户搜索历史和风格偏好，无法做个性化推荐。', indent=True)

h2(doc, '反思')
p(doc, '这次项目经历了一个完整的"试错→收敛"过程。最初版本是多智能体宿舍模拟器——四个AI室友拥有MBTI人格，互相对话产生好感度变化，架构复杂但未解决任何实际问题。第二版改成四步诊断式期末突击教练——功能完整但流程过长，演示容易翻车。最终收敛为"拾句"——极其简单，但因简单而稳定。', indent=True)
p(doc, '从中获得几点体会：', indent=True, bold=True)
p(doc, '第一，AI产品的价值不在于用了多少技术，而在于解决了什么真实问题。多智能体宿舍用了大量Prompt工程和状态管理，但没有用户真正需要它。拾句的技术难度远低于前者，但因为它解决了"想表达但找不到合适语言"这个真实需求，反而更有价值。', indent=True)
p(doc, '第二，最难的不是调用API，而是想清楚AI应该做什么、不应该做什么。拾句的核心决策是：让AI找诗而不是写诗。这一举解决了内容质量问题——展示的是李白杜甫的原句，不是AI的即兴创作。"限制AI"有时比"发挥AI"更有产品价值。', indent=True)
p(doc, '第三，演示稳定性在评审场景中比功能完整性更重要。评委只花3分钟看项目，一个功能简单但点击必出结果的产品，比功能复杂但可能卡顿的产品更容易拿高分。', indent=True)
p(doc, '第四，场景共鸣是最大的杠杆。大学生正在经历的几乎所有情感——孤独、迷茫、暗恋、思乡、毕业焦虑——都被古人写透了。"用古典文学回应现代情绪"这个概念，任何文化背景的人都能一秒理解并产生共鸣。', indent=True)

# ── 六 (with TABLE) ──
h1(doc, '六、交付物清单')
p(doc, '')

table = doc.add_table(rows=6, cols=2, style='Light Grid Accent 1')
table.autofit = True
items = [
    ('交付项', '内容 / 链接'),
    ('智能体线上地址', 'https://ai-agent-work-as2jgjymz7emt2w5avocqv.streamlit.app'),
    ('GitHub 代码仓库', 'https://github.com/ZhiRam/AI-agent-work'),
    ('3分钟讲解视频', '见单独上传的视频文件'),
    ('AI交互记录', '见视频内演示截图及录屏'),
    ('本报告（Word版）', '桌面\\word\\拾句_作业报告.docx'),
]
for i, (k, v) in enumerate(items):
    row = table.rows[i]
    row.cells[0].text = k
    row.cells[1].text = v
    for cell in [row.cells[0], row.cells[1]]:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.name = '微软雅黑'
                run.font.size = Pt(10)
        if i == 0:
            for para in row.cells[0].paragraphs:
                for run in para.runs: run.bold = True
            for para in row.cells[1].paragraphs:
                for run in para.runs: run.bold = True

doc.add_paragraph()
sep(doc)
p(doc, '全文完', size=10, center=True, color=GREEN_MUTED)
doc.save(r'C:/Users/ZhiRan/Desktop/word/report.docx')
print('REPORT DONE')

# ═══════════════════════════════════════
# FILE 2: SCRIPT (academic)
# ═══════════════════════════════════════
doc2 = setup_doc()

# COVER
for _ in range(6): doc2.add_paragraph()
p(doc2, '拾  句', bold=True, size=30, center=True, color=GREEN)
p(doc2, '演示讲解脚本', size=18, center=True, color=GRAY)
for _ in range(3): doc2.add_paragraph()
p(doc2, '王  宸    2025212714', size=14, center=True, color=DARK)
p(doc2, '北京邮电大学 · AI智能体开发', size=11, center=True, color=GRAY)
doc2.add_page_break()

h1(doc2, 'Slide 1：封面（约5秒）')
p(doc2, '各位老师、同学大家好。我今天展示的项目名为「拾句」——一词入，千年应。本项目核心定位为：基于大语言模型的古典文学语义匹配工具，通过Prompt工程实现对唐诗宋词等古典文本的精准检索与意境化呈现。', indent=True)

h1(doc2, 'Slide 2：项目概述（约50秒）')
p(doc2, '首先阐述项目的出发点。在日常社交媒体的表达场景中，个体常常面临"情感丰富但语言匮乏"的困境。而中国古典文学——从唐诗宋词到明清散文——实际上已经以极高的文学完成度覆盖了人类情感的几乎全部光谱。问题不在于缺乏好的表达，而在于个体与古典文本之间存在信息检索的鸿沟。', indent=True)
p(doc2, '本项目的核心功能可概括为：用户输入一个关键词，系统从古典文学语料中匹配最贴切的名句，并以视觉卡片的形式呈现。需要特别指出的是，本项目的设计哲学与当前主流的AI内容生成范式存在根本差异——我们拒绝让AI进行文学创作，而是利用大语言模型在训练过程中内化的古典文学知识，实现语义层面的"检索式匹配"。简言之，不是让AI写诗，而是让AI帮助用户找到已有的好诗。', indent=True)
p(doc2, '系统提供四种文学体裁风格供用户切换：唐诗气象、宋词余韵、古文雅意、散文光阴。同一关键词在不同体裁下将获得差异化的文学回应，构成"一词四境"的体验模式。', indent=True)

h1(doc2, 'Slide 3：技术架构（约30秒）')
p(doc2, '在技术实现层面，本项目采用Python语言开发，前端基于Streamlit框架构建，大语言模型层调用DeepSeek-V3，通过硅基流动API平台进行代理访问，兼容OpenAI SDK格式。', indent=True)
p(doc2, '核心的技术工作集中在Prompt工程层面。四种文学风格各自对应一套独立的System Prompt，明确指定了作者范围、文风特征和输出格式约束。输出采用结构化JSON格式，包含quote（引用原文）、source（作者与篇名）、paraphrase（白话解读）、emotion（情绪标签）四个字段。其中source字段的强制校验机制有效地抑制了大语言模型的幻觉问题。部署方面采用Streamlit Cloud免费托管方案，与GitHub仓库自动同步，敏感信息通过Secrets机制管理。', indent=True)

h1(doc2, 'Slide 4：核心创新（约50秒）')
p(doc2, '本项目的创新点可归纳为四个方面。', indent=True)
p(doc2, '第一，创作范式的逆向设计。当前AI应用的主流趋势是"生成式"——让模型创造内容。本项目反其道而行之，采用"检索式"策略，利用模型的语义理解能力在训练数据中的古典文学知识库中进行匹配。这避免了生成式模型在文学创作中普遍存在的质量低下和幻觉问题，同时充分发挥了古典文学本身的审美价值。', indent=True)
p(doc2, '第二，一文多境的体验设计。同一关键词在四种文学体裁中呈现出差异化的美学表达。以"孤独"为例：唐诗提供的是"前不见古人，后不见来者"的苍茫辽阔；宋词呈现的是"无可奈何花落去，似曾相识燕归来"的婉约低徊；古文展现的是"庭有枇杷树，今已亭亭如盖矣"的深情克制；散文则是"热闹是他们的，我什么也没有"的现代性孤独。四种风格构成了一次横跨千年的文学巡游。', indent=True)
p(doc2, '第三，交互极简主义。整个用户界面仅包含三个可操作元素，无需用户编写Prompt或调整任何参数。从输入到获得结果的平均操作时间不超过3秒，显著降低了AI工具的使用门槛。生成的卡片兼具审美价值与传播属性，实现了"产品即内容"的设计理念。', indent=True)
p(doc2, '第四，视觉符号与文学传统的一一对应。唐诗风格采用宣纸基色、楷体字栈与朱红圆形印章，呼应盛唐气象；宋词风格采用淡紫渐变与柔光圆环，对应婉约意境；古文风格以古纸色调与虚线邮戳印记呈现文人的雅致趣味；散文风格则使用暖黄底色与衬线字体，传递人间烟火的温度。视觉语言在此不是独立的装饰层，而是文学体验的构成性要素。', indent=True)

h1(doc2, 'Slide 5：问题与反思（约35秒）')
p(doc2, '在开发过程中，遇到若干具有代表性的技术问题。其一，大语言模型的幻觉问题可能导致其编造不存在的古典诗句，解决方案为在Prompt中限定作者范围并要求输出标注来源；其二，项目部署于Streamlit Cloud（美国服务器），DeepSeek官方API拒绝海外IP访问，通过切换至国内代理平台硅基流动得以解决；其三，CSS与HTML在Streamlit中的渲染兼容性问题，通过f-string构建与html.escape()转义进行了修复。', indent=True)
p(doc2, '本项目经历了三次重大版本迭代。初版为多智能体宿舍社交模拟器，架构复杂度较高但缺乏实际应用价值。第二版转为四阶段诊断式期末复习教练，功能完整但操作流程冗长，在演示场景中稳定性不足。第三版即当前版本——大幅简化功能范围，聚焦单一核心能力，以稳定性换取演示的可靠性。', indent=True)
p(doc2, '从这一迭代过程中获得了以下认识：第一，人工智能应用的产品价值取决于其解决的真实问题，而非技术复杂度；第二，在工程实践中，明确模型的"不应为"往往比拓展其"可为"更具挑战和意义；第三，在评审性质的演示场景中，系统的可靠性和稳定性优先于功能的完整性；第四，文化共鸣是产品接受度的核心杠杆——"以古典文学回应现代情感"这一概念框架具有跨背景的可理解性。', indent=True)

h1(doc2, 'Slide 6：致谢（约10秒）')
p(doc2, '以上为项目的全部介绍。线上体验地址已在交付物清单中列出，欢迎各位老师和同学体验并提供反馈。谢谢。', indent=True)

h1(doc2, '演示操作配合')
p(doc2, '在讲解至Slide 4（核心创新）时，建议同步进行浏览器端的实际操作演示：打开应用链接，输入示例关键词"孤独"；选择"唐诗气象"风格，生成卡片并展示效果；切换至"宋词余韵"风格，展示同一关键词在不同文学体裁下的差异化输出；说明卡片可截图并直接用于社交媒体分享。实际操作时长约1分钟，整体演示严格控制在3分钟以内。', indent=True)

doc2.add_page_break()

# ═══════════════════════════════════════
# PART 2: 专属演示视频录制脚本
# ═══════════════════════════════════════
p(doc2, '专属演示视频录制脚本', bold=True, size=22, center=True, color=GREEN)
p(doc2, '（以下为视频录制的逐秒操作指南，可直接照着录）', size=12, center=True, color=GRAY)
doc2.add_paragraph()

# ── 片头 ──
h1(doc2, '一、片头（0:00 - 0:15）')
p(doc2, '画面：', bold=True)
p(doc2, '屏幕干净，关闭无关窗口和通知。打开项目线上链接，显示首页（输入框 + 四种风格按钮）。标题「拾句」可见。')
p(doc2, '')
p(doc2, '旁白（对着麦克风说）：', bold=True)
p(doc2, '"大家好，我今天展示的项目叫「拾句」。它的功能很简单：你输入一个词，它从唐诗宋词里帮你找到那句最贴切的古典名句，生成一张诗意卡片。"', indent=True)

# ── 演示 1 ──
h1(doc2, '二、第一次演示：唐诗气象（0:15 - 0:55）')
p(doc2, '操作步骤：', bold=True)
p(doc2, '1. 鼠标移到输入框，慢慢输入"孤独"两个字，让观众看见')
p(doc2, '2. 鼠标指向风格按钮，当前选中"唐诗气象"（保持默认）')
p(doc2, '3. 点击「生成」按钮')
p(doc2, '4. 等待卡片出现（约1-2秒），鼠标在卡片上缓慢划过，让观众看清内容')
p(doc2, '')
p(doc2, '旁白：', bold=True)
p(doc2, '"我输入一个词——孤独。选唐诗气象，点生成。"', indent=True)
p(doc2, '等卡片出来——')
p(doc2, '"你看，陈子昂的《登幽州台歌》。前不见古人，后不见来者，念天地之悠悠，独怆然而涕下。千年前一个人站在天地之间的孤独，和今天你深夜发朋友圈的孤独，是同一种。这就是唐诗的苍茫。"', indent=True)
p(doc2, '')
p(doc2, '注意：', bold=True)
p(doc2, '卡片出来后停留5秒，镜头不要晃，给评委看清的时间。')

# ── 演示 2 ──
h1(doc2, '三、第二次演示：切换风格（0:55 - 1:25）')
p(doc2, '操作步骤：', bold=True)
p(doc2, '1. 鼠标点击"宋词余韵"风格按钮')
p(doc2, '2. 再次点击「生成」')
p(doc2, '3. 新的卡片出现，内容和风格完全不同')
p(doc2, '')
p(doc2, '旁白：', bold=True)
p(doc2, '"现在换到宋词。同一个词——孤独。李清照说，此情无计可消除，才下眉头，却上心头。完全没有提到孤独两个字，但每个字都在说孤独。这就是一词四境——同一个情绪，在四种文学传统里，得到四种完全不同的回应。"', indent=True)
p(doc2, '')
p(doc2, '技巧：', bold=True)
p(doc2, '为了让两张卡片形成对比，可以用鼠标在两页之间快速切换（Streamlit没有后退按钮，建议提前准备好两张卡的截图，录屏时叠放展示）。')

# ── 演示 3 ──
h1(doc2, '四、展示分享场景（1:25 - 1:45）')
p(doc2, '操作步骤：', bold=True)
p(doc2, '1. 对卡片截图（或用系统截图工具框选卡片区域）')
p(doc2, '2. 打开微信/朋友圈界面（提前准备好空白朋友圈编辑页）')
p(doc2, '3. 把截图粘贴进去，配上文字"深夜emo被古人治愈了"')
p(doc2, '')
p(doc2, '旁白：', bold=True)
p(doc2, '"而且你发现没有——这张卡片设计得很精致。这是唐诗风格，宣纸底色、楷体字、红色印章，像一幅小画。直接截图发朋友圈，不用修图。产品即内容。"', indent=True)
p(doc2, '')
p(doc2, '提示：', bold=True)
p(doc2, '如果嫌微信界面切换麻烦，可以直接用PPT叠一张手机框的mockup图，把卡片截图放进去。')

# ── 技术亮点 ──
h1(doc2, '五、技术亮点（1:45 - 2:20）')
p(doc2, '画面：', bold=True)
p(doc2, '切到PPT的技术架构页（Slide 3），指一下关键层级。')
p(doc2, '')
p(doc2, '旁白：', bold=True)
p(doc2, '"技术上，核心是Prompt工程。四种风格各自一套System Prompt，输出JSON结构化格式。为什么是JSON？因为要保证卡片内容稳定——quote字段是引用原文，source字段强制标注作者和篇名，杜绝了AI编造假诗的可能。部署在Streamlit Cloud，免费托管，GitHub一推自动更新。"', indent=True)
p(doc2, '')
p(doc2, '注意：', bold=True)
p(doc2, '技术部分语速可以稍快，控制在35秒以内。评委听不懂细节没关系，关键是听到几个关键词：Prompt工程、JSON结构化、杜绝幻觉。')

# ── 收尾 ──
h1(doc2, '六、收尾（2:20 - 2:45）')
p(doc2, '画面：', bold=True)
p(doc2, '切回项目页面，鼠标在四种风格按钮上依次悬停（唐诗、宋词、古文、散文），最后停在输入框。')
p(doc2, '')
p(doc2, '旁白：', bold=True)
p(doc2, '"这个项目改了三版。第一版多智能体宿舍，第二版期末教练，最后收敛到这一版——拾句。极其简单，但因为简单所以稳定，因为稳定所以能用。"', indent=True)
p(doc2, '"谢谢大家。链接在简介里，欢迎体验。"', indent=True)

# ── 技术 Checklist ──
h1(doc2, '七、录制前检查清单')
dot(doc2, '关闭微信、QQ 等桌面通知，避免录屏时弹出消息')
dot(doc2, '浏览器开无痕模式，避免地址栏出现其他历史记录')
dot(doc2, '提前打开项目链接，确认页面加载正常、API 可用')
dot(doc2, '准备两个测试词（"孤独"+"暗恋"或"故乡"），提前跑一遍确认不卡')
dot(doc2, '麦克风测试：用手机录一段听一下，确保声音清晰无杂音')
dot(doc2, '如果可能，用 OBS 录屏（免费），画质比系统自带的好')
dot(doc2, '总时长控制在 2分45秒 - 3分钟，不超时')

doc2.add_paragraph()
sep(doc2)
p(doc2, '全文完', size=10, center=True, color=GREEN_MUTED)
doc2.save(r'C:/Users/ZhiRan/Desktop/ppt/script.docx')
print('SCRIPT DONE')
